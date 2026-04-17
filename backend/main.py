"""DocuGen Open local backend."""
from __future__ import annotations

import io
import json
import logging
import re
import subprocess
from urllib.parse import quote, urlparse
from typing import List, Optional

import httpx
from fastapi import FastAPI, File, Form, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, field_validator

from pdf_bridge import DOCX_MEDIA_TYPE, export_docx_with_offline_pdf, extract_with_offline_pdf

log = logging.getLogger("docugen")
ALLOWED_LOCAL_HOSTS = {"127.0.0.1", "localhost", "::1"}
IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tif", ".tiff")

CHUNK_LIMIT = 4000
BID_CHUNK_LIMIT = 12000
MAX_FILE_SIZE = 10 * 1024 * 1024
BID_MAX_FILE_SIZE = 50 * 1024 * 1024

app = FastAPI(title="DocuGen Open Backend", version="0.2.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://127.0.0.1:5173",
        "http://localhost:5173",
        "http://127.0.0.1:9000",
        "http://localhost:9000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class LocalAIConfig(BaseModel):
    apiBaseUrl: str
    apiKey: str = ""
    model: str
    bidModel: Optional[str] = None

    @field_validator("apiBaseUrl")
    @classmethod
    def normalize_base_url(cls, value: str) -> str:
        value = value.strip().rstrip("/")
        if not value:
            raise ValueError("API Base URL 不能为空")

        parsed = urlparse(value)
        if parsed.scheme != "https":
            if parsed.scheme != "http" or parsed.hostname not in ALLOWED_LOCAL_HOSTS:
                raise ValueError("API Base URL 仅支持 HTTPS，或本机 HTTP 地址")

        if not parsed.netloc:
            raise ValueError("API Base URL 格式无效")
        return value

    @field_validator("model")
    @classmethod
    def require_model(cls, value: str) -> str:
        value = value.strip()
        if not value:
            raise ValueError("默认模型不能为空")
        return value

    @field_validator("apiKey")
    @classmethod
    def normalize_api_key(cls, value: str) -> str:
        return value.strip()

    @field_validator("bidModel")
    @classmethod
    def normalize_bid_model(cls, value: Optional[str]) -> Optional[str]:
        if value is None:
            return None
        value = value.strip()
        return value or None


class ChatRequest(BaseModel):
    prompt: str
    model: Optional[str] = None
    config: LocalAIConfig


class ChunkProcessRequest(BaseModel):
    text: str
    mode: str = "format"
    model: Optional[str] = None
    config: LocalAIConfig


def _parse_config_json(raw: str) -> LocalAIConfig:
    try:
        return LocalAIConfig.model_validate(json.loads(raw))
    except Exception as exc:
        raise ValueError(f"配置无效: {exc}") from exc


def _safe_upstream_error(exc: httpx.HTTPStatusError) -> str:
    status_code = exc.response.status_code if exc.response is not None else 502
    return f"上游模型服务返回错误 (HTTP {status_code})"


def _size_error(uploaded_size: int, size_limit: int) -> Optional[JSONResponse]:
    if uploaded_size > size_limit:
        limit_mb = size_limit // (1024 * 1024)
        return JSONResponse(status_code=413, content={"error": f"文件超过 {limit_mb}MB 限制"})
    return None


def _read_docx(content: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))
    parts: list[str] = []

    for para in doc.paragraphs:
      if para.text:
        parts.append(para.text)

    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            header = rows[0]
            separator = " | ".join(["---"] * len(table.rows[0].cells))
            parts.append("\n" + "\n".join([header, separator, *rows[1:]]) + "\n")

    return "\n".join(parts).strip()


def _read_excel(content: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for worksheet in workbook.worksheets:
            parts.append(f"## {worksheet.title}\n")
            rows_data: list[list[str]] = []
            for row in worksheet.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                if any(cells):
                    rows_data.append(cells)
            if rows_data:
                header = " | ".join(rows_data[0])
                separator = " | ".join(["---"] * len(rows_data[0]))
                body = "\n".join(" | ".join(row) for row in rows_data[1:])
                parts.append(f"{header}\n{separator}\n{body}\n")
    finally:
        workbook.close()

    return "\n".join(parts).strip()


def _extract_pdf_or_image(content: bytes, filename: str) -> tuple[str, dict]:
    payload = extract_with_offline_pdf(content, filename)
    if payload and payload.get("text"):
        text = str(payload["text"]).strip()
        if text:
            return text, payload.get("meta") or {}

    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        try:
            result = subprocess.run(
                ["pdftotext", "-layout", "-", "-"],
                input=content,
                capture_output=True,
                timeout=30,
            )
        except FileNotFoundError:
            return "", {}
        except Exception as exc:
            log.warning("pdftotext fallback failed: %s", exc)
            return "", {}

        if result.returncode == 0:
            return (
                result.stdout.decode("utf-8", errors="replace").strip(),
                {"summary": "pdftotext fallback"},
            )
    return "", {}


def _decode_plain_text(content: bytes) -> str:
    for encoding in ["utf-8-sig", "utf-8", "gbk", "gb2312", "gb18030", "latin-1"]:
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return ""


async def _call_ai(prompt: str, model: str, config: LocalAIConfig) -> str:
    headers = {"Content-Type": "application/json"}
    if config.apiKey:
        headers["Authorization"] = f"Bearer {config.apiKey}"

    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
    }

    async with httpx.AsyncClient(timeout=180.0) as client:
        response = await client.post(
            f"{config.apiBaseUrl}/chat/completions",
            json=payload,
            headers=headers,
        )
        response.raise_for_status()
        data = response.json()
        return (
            data.get("choices", [{}])[0]
            .get("message", {})
            .get("content", "")
            .strip()
        )


def _build_format_prompt(text: str, part_info: str = "") -> str:
    return f"""你是一名专业的文档格式整理助手。请把下面的内容整理成结构清晰、标题层级合理的 Markdown 文档。

核心规则:
1. 绝对不要删减、改写、总结、补充原文内容。
2. 只允许调整结构、标题层级、段落、列表和表格形式。
3. 识别出表格时，请输出为 Markdown 表格。
4. 只返回 Markdown 内容，不要额外解释，不要包裹代码块。
{part_info}

原始内容:
{text}"""


def _build_table_prompt(text: str, part_info: str = "") -> str:
    return f"""你是一名专业的文档表格整理助手。请分析下面的内容，把适合表格化的信息整理成 Markdown 表格。

核心规则:
1. 不要删减、改写、总结、补充原文内容。
2. 适合转表格的内容请使用 Markdown 表格表示。
3. 不适合转表格的说明文字保留在表格外。
4. 只返回 Markdown 内容，不要额外解释，不要包裹代码块。
{part_info}

原始内容:
{text}"""


def _build_bid_prompt(text: str, part_info: str = "") -> str:
    return f"""你是一名专业的标书格式整理助手。请把下面的标书内容整理成规范的 Markdown 文档。

严格要求:
1. 不要删减、改写、总结、补充任何原文内容。
2. 只允许调整标题、段落、编号、列表和表格形式。
3. 表格信息请尽量保持完整。
4. 只返回 Markdown 内容，不要额外解释，不要包裹代码块。
{part_info}

原始内容:
{text}"""


def _split_into_chunks(text: str, limit: int = CHUNK_LIMIT) -> List[str]:
    paragraphs = re.split(r"\n\s*\n", text)
    chunks: List[str] = []
    current = ""

    for para in paragraphs:
        if len(current) + len(para) + 2 > limit and current:
            chunks.append(current)
            current = para
        else:
            current = f"{current}\n\n{para}" if current else para

    if current:
        chunks.append(current)

    final: List[str] = []
    for chunk in chunks:
        if len(chunk) <= limit:
            final.append(chunk)
            continue

        lines = chunk.split("\n")
        buffer = ""
        for line in lines:
            if len(buffer) + len(line) + 1 > limit and buffer:
                final.append(buffer)
                buffer = line
            else:
                buffer = f"{buffer}\n{line}" if buffer else line
        if buffer:
            final.append(buffer)

    return final or [text]


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/api/models")
async def list_models(config: LocalAIConfig):
    headers = {"Authorization": f"Bearer {config.apiKey}"} if config.apiKey else {}
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.get(f"{config.apiBaseUrl}/models", headers=headers)
            response.raise_for_status()
            data = response.json()
    except Exception as exc:
        return {
            "error": f"模型探测失败: {exc}",
            "models": [{"id": config.model, "name": config.model}],
            "default": config.model,
        }

    items = []
    for item in data.get("data", []):
        if isinstance(item, dict) and item.get("id"):
            model_id = item["id"]
            items.append({"id": model_id, "name": model_id})

    if not items:
        items = [{"id": config.model, "name": config.model}]

    return {"models": items, "default": config.model}


@app.post("/api/ai/chat")
async def ai_chat(req: ChatRequest):
    model = req.model or req.config.model
    try:
        text = await _call_ai(req.prompt, model, req.config)
        return {"text": text, "model": model}
    except httpx.HTTPStatusError as exc:
        return JSONResponse(status_code=exc.response.status_code, content={"error": _safe_upstream_error(exc)})
    except Exception as exc:
        return JSONResponse(status_code=500, content={"error": f"AI 调用异常: {exc}"})


@app.post("/api/ai/chunk-process")
async def chunk_process(req: ChunkProcessRequest):
    text = req.text
    is_bid = req.mode == "bid"
    model = req.config.bidModel if is_bid and req.config.bidModel else (req.model or req.config.model)
    chunk_limit = BID_CHUNK_LIMIT if is_bid else CHUNK_LIMIT
    build_prompt = _build_bid_prompt if is_bid else (_build_format_prompt if req.mode == "format" else _build_table_prompt)

    if len(text) <= chunk_limit:
        try:
            result = await _call_ai(build_prompt(text), model, req.config)
            return {"text": result, "chunks_processed": 1, "model": model}
        except httpx.HTTPStatusError as exc:
            return JSONResponse(status_code=exc.response.status_code, content={"error": _safe_upstream_error(exc)})
        except Exception as exc:
            return JSONResponse(status_code=500, content={"error": str(exc)})

    chunks = _split_into_chunks(text, limit=chunk_limit)
    total = len(chunks)
    results: List[str] = []

    for index, chunk in enumerate(chunks, 1):
        if is_bid:
            part_info = f"\n注意: 这是完整文档的第 {index}/{total} 部分，请保持与前后部分风格一致，不要省略内容。"
        else:
            part_info = f"\n注意: 这是完整文档的第 {index}/{total} 部分，请保持风格一致，不要添加总结。"

        try:
            result = await _call_ai(build_prompt(chunk, part_info), model, req.config)
            results.append(result)
        except Exception as exc:
            log.warning("Chunk %s/%s failed: %s", index, total, exc)
            results.append(chunk)

    return {"text": "\n\n".join(results), "chunks_processed": total, "model": model}


@app.post("/api/upload")
async def upload_file(
    file: UploadFile = File(...),
    config: str = Form(...),
    bid: bool = False,
):
    try:
        _parse_config_json(config)
    except ValueError as exc:
        return JSONResponse(status_code=400, content={"error": str(exc)})

    size_limit = BID_MAX_FILE_SIZE if bid else MAX_FILE_SIZE
    content = await file.read(size_limit + 1)
    size_error = _size_error(len(content), size_limit)
    if size_error:
        return size_error

    filename = file.filename or ""
    lower_name = filename.lower()
    text = ""
    meta: dict = {}

    try:
        if lower_name.endswith(".docx"):
            text = _read_docx(content)
        elif lower_name.endswith(".doc"):
            return JSONResponse(status_code=400, content={"error": "暂不支持 .doc，请先另存为 .docx 再导入。"})
        elif lower_name.endswith((".xlsx", ".xls")):
            text = _read_excel(content)
        elif lower_name.endswith(".pdf") or lower_name.endswith(IMAGE_EXTENSIONS):
            text, meta = _extract_pdf_or_image(content, filename)
            if not text and lower_name.endswith(IMAGE_EXTENSIONS):
                return JSONResponse(
                    status_code=400,
                    content={"error": "图片解析需要配置本地离线引擎。请设置 DOCUGEN_OFFLINE_PDF_ROOT 指向 Offline_PDF_Structure。"},
                )
        else:
            text = _decode_plain_text(content)
    except Exception as exc:
        return JSONResponse(status_code=400, content={"error": f"文件解析失败: {exc}"})

    if not text:
        return JSONResponse(status_code=400, content={"error": "无法识别文件内容或编码"})

    return {"text": text, "filename": filename, "chars": len(text), "meta": meta}


@app.post("/api/export/pdf-docx")
async def export_pdf_docx(file: UploadFile = File(...)):
    filename = file.filename or ""
    if not filename.lower().endswith(".pdf"):
        return JSONResponse(status_code=400, content={"error": "请上传 PDF 文件。"})

    content = await file.read(BID_MAX_FILE_SIZE + 1)
    size_error = _size_error(len(content), BID_MAX_FILE_SIZE)
    if size_error:
        return size_error

    output_bytes, download_name, meta = export_docx_with_offline_pdf(content, filename)
    if not output_bytes or not download_name:
        return JSONResponse(
            status_code=503,
            content={
                "error": "未找到可用的本地 PDF 转 Word 引擎。请设置 DOCUGEN_OFFLINE_PDF_ROOT 指向 Offline_PDF_Structure，或把离线包放到 backend/offline_pdf_bundle。",
            },
        )

    encoded_name = quote(download_name)
    headers = {
        "Content-Disposition": f"attachment; filename=\"{download_name}\"; filename*=UTF-8''{encoded_name}",
    }
    if meta.get("summary"):
        headers["X-DocuGen-Summary"] = str(meta["summary"])

    return StreamingResponse(io.BytesIO(output_bytes), media_type=DOCX_MEDIA_TYPE, headers=headers)

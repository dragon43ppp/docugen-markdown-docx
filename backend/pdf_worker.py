from __future__ import annotations

import argparse
import io
import json
import os
import sys
import warnings
from pathlib import Path
from typing import Iterable

warnings.filterwarnings("ignore")
os.environ.setdefault("DISABLE_MODEL_SOURCE_CHECK", "True")

import camelot
import cv2
import fitz
import numpy as np
import pandas as pd
from docx import Document
from pdf2docx import Converter
from paddleocr import PPStructure, PaddleOCR

ENGINE_CACHE: dict[tuple[str, float], tuple[PPStructure | None, PaddleOCR | None]] = {}


def normalize_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""
    return str(value).replace("\xa0", " ").strip()


def escape_markdown_cell(value: object) -> str:
    text = normalize_text(value)
    return text.replace("|", r"\|").replace("\r\n", "\n").replace("\r", "\n").replace("\n", "<br>")


def dataframe_to_markdown(df: pd.DataFrame) -> str:
    if df.empty and len(df.columns) == 0:
        return ""

    header = [escape_markdown_cell(col) for col in df.columns.tolist()]
    if not header:
        header = [f"Column {index + 1}" for index in range(max(df.shape[1], 1))]

    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * len(header)) + " |",
    ]

    if df.shape[0] == 0:
        return "\n".join(lines)

    for row in df.itertuples(index=False):
        cells = [escape_markdown_cell(cell) for cell in row]
        if len(cells) < len(header):
            cells.extend([""] * (len(header) - len(cells)))
        lines.append("| " + " | ".join(cells[: len(header)]) + " |")
    return "\n".join(lines)


def html_to_markdown_table(html: str) -> str:
    frames = pd.read_html(io.StringIO(html))
    if not frames:
        return ""
    return dataframe_to_markdown(frames[0].fillna(""))


def extract_block_lines(payload: object) -> list[str]:
    if isinstance(payload, str):
        text = normalize_text(payload)
        return [text] if text else []

    lines: list[str] = []
    if isinstance(payload, list):
        for item in payload:
            if isinstance(item, dict):
                text = normalize_text(item.get("text") or item.get("res"))
                if text:
                    lines.append(text)
            else:
                text = normalize_text(item)
                if text:
                    lines.append(text)
    return lines


def ppstructure_to_markdown(blocks: object) -> tuple[str, int, int]:
    if not isinstance(blocks, list):
        return "", 0, 0

    parts: list[str] = []
    table_count = 0
    char_count = 0
    sorted_blocks = sorted(
        (block for block in blocks if isinstance(block, dict)),
        key=lambda block: (
            block.get("bbox", [0, 0, 0, 0])[1],
            block.get("bbox", [0, 0, 0, 0])[0],
        ),
    )

    for block in sorted_blocks:
        block_type = normalize_text(block.get("type")).lower()
        if block_type == "table":
            html = ""
            if isinstance(block.get("res"), dict):
                html = normalize_text(block["res"].get("html"))
            if not html:
                continue
            try:
                table_md = html_to_markdown_table(html)
            except Exception:
                continue
            if table_md:
                parts.append(table_md)
                table_count += 1
                char_count += len(table_md)
            continue

        lines = extract_block_lines(block.get("res"))
        if not lines:
            continue

        if block_type == "title":
            text = "## " + normalize_text(" ".join(lines))
        elif block_type == "list":
            text = "\n".join(f"- {line}" for line in lines if normalize_text(line))
        else:
            text = "\n".join(lines)
        text = text.strip()
        if not text:
            continue
        parts.append(text)
        char_count += len(text)

    return "\n\n".join(parts).strip(), table_count, char_count


def ocr_to_markdown(raw: object) -> tuple[str, int]:
    current = raw
    if isinstance(raw, list) and len(raw) == 1 and isinstance(raw[0], list):
        current = raw[0]

    lines: list[str] = []
    if isinstance(current, list):
        for item in current:
            if not isinstance(item, list) or len(item) < 2:
                continue
            payload = item[1]
            if isinstance(payload, (list, tuple)) and payload:
                text = normalize_text(payload[0])
            else:
                text = normalize_text(payload)
            if text:
                lines.append(text)

    text = "\n".join(lines).strip()
    return text, len(text)


def render_page_bgr(page: fitz.Page) -> np.ndarray:
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
    if pix.n == 4:
        img = cv2.cvtColor(img, cv2.COLOR_RGBA2RGB)
    return cv2.cvtColor(img, cv2.COLOR_RGB2BGR)


def is_scanned_page(page: fitz.Page) -> bool:
    text = page.get_text("text").strip()
    if len(text) > 80:
        return False

    pix = page.get_pixmap(dpi=72)
    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
    if pix.n == 4:
        img = cv2.cvtColor(img, cv2.COLOR_RGBA2RGB)
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    edges = cv2.Canny(gray, 50, 150)
    edge_density = float(np.mean(edges))
    if edge_density < 2.0:
        return True
    return len(text) < 60


def get_engines(model_dir: Path | None, threshold: float) -> tuple[PPStructure | None, PaddleOCR | None]:
    if model_dir is None or not model_dir.exists():
        return None, None

    key = (str(model_dir), threshold)
    cached = ENGINE_CACHE.get(key)
    if cached is not None:
        return cached

    dict_path = model_dir / "ppocr_keys_v1.txt"
    required = [
        dict_path,
        model_dir / "picodet_lcnet_x1_0_fgd_layout_cdla_infer",
        model_dir / "ch_ppstructure_mobile_v2.0_SLANet_infer",
        model_dir / "ch_PP-OCRv4_det_infer",
        model_dir / "ch_PP-OCRv4_rec_infer",
        model_dir / "ch_ppocr_mobile_v2.0_cls_infer",
    ]
    if not all(path.exists() for path in required):
        ENGINE_CACHE[key] = (None, None)
        return ENGINE_CACHE[key]

    try:
        struct_engine = PPStructure(
            layout_model_dir=str(model_dir / "picodet_lcnet_x1_0_fgd_layout_cdla_infer"),
            table_model_dir=str(model_dir / "ch_ppstructure_mobile_v2.0_SLANet_infer"),
            det_model_dir=str(model_dir / "ch_PP-OCRv4_det_infer"),
            rec_model_dir=str(model_dir / "ch_PP-OCRv4_rec_infer"),
            rec_char_dict_path=str(dict_path),
            use_gpu=False,
            show_log=False,
            table=True,
            layout=True,
            structure_threshold=threshold,
        )
        ocr_engine = PaddleOCR(
            det_model_dir=str(model_dir / "ch_PP-OCRv4_det_infer"),
            rec_model_dir=str(model_dir / "ch_PP-OCRv4_rec_infer"),
            cls_model_dir=str(model_dir / "ch_ppocr_mobile_v2.0_cls_infer"),
            rec_char_dict_path=str(dict_path),
            use_angle_cls=True,
            use_gpu=False,
            show_log=False,
        )
        ENGINE_CACHE[key] = (struct_engine, ocr_engine)
    except Exception:
        ENGINE_CACHE[key] = (None, None)
    return ENGINE_CACHE[key]


def extract_native_tables(pdf_path: Path, page_number: int) -> tuple[list[pd.DataFrame], str | None]:
    for flavor in ("stream", "lattice"):
        try:
            tables = camelot.read_pdf(str(pdf_path), pages=str(page_number), flavor=flavor)
        except Exception:
            continue
        if len(tables) > 0:
            return [table.df.fillna("") for table in tables], flavor
    return [], None


def add_page_heading(parts: list[str], page_number: int, total_pages: int, content: str) -> None:
    content = content.strip()
    if not content:
        return
    if total_pages > 1:
        parts.append(f"## Page {page_number}\n\n{content}")
    else:
        parts.append(content)


def build_summary(
    kind: str,
    total_pages: int,
    scanned_pages: int,
    native_pages: int,
    table_count: int,
    ocr_fallback_pages: int,
    mode: str,
) -> str:
    summary_bits = [mode, f"pages: {total_pages}"]
    if scanned_pages:
        summary_bits.append(f"scanned: {scanned_pages}")
    if native_pages:
        summary_bits.append(f"native: {native_pages}")
    if table_count:
        summary_bits.append(f"tables: {table_count}")
    if ocr_fallback_pages:
        summary_bits.append(f"ocr-fallback: {ocr_fallback_pages}")
    if kind == "image":
        summary_bits = [mode, kind] + [bit for bit in summary_bits[1:] if not bit.startswith("pages:")]
    return " | ".join(summary_bits)


def process_pdf(input_path: Path, model_dir: Path | None, threshold: float) -> dict:
    struct_engine, ocr_engine = get_engines(model_dir, threshold)
    doc = fitz.open(str(input_path))
    total_pages = len(doc)
    parts: list[str] = []
    scanned_pages = 0
    native_pages = 0
    table_count = 0
    ocr_fallback_pages = 0
    camelot_pages = 0

    for index in range(total_pages):
        page = doc.load_page(index)
        page_number = index + 1
        if is_scanned_page(page):
            scanned_pages += 1
            page_text = page.get_text("text").strip()
            if struct_engine is not None or ocr_engine is not None:
                image = render_page_bgr(page)
                structured_md = ""
                structured_chars = 0
                structured_tables = 0
                if struct_engine is not None:
                    structured_md, structured_tables, structured_chars = ppstructure_to_markdown(struct_engine(image))
                    table_count += structured_tables

                if ocr_engine is not None and structured_chars < 80:
                    ocr_md, ocr_chars = ocr_to_markdown(ocr_engine.ocr(image, cls=True))
                    if ocr_chars > structured_chars:
                        structured_md = ocr_md
                        ocr_fallback_pages += 1

                add_page_heading(parts, page_number, total_pages, structured_md or page_text)
            else:
                add_page_heading(parts, page_number, total_pages, page_text)
            continue

        native_pages += 1
        page_parts: list[str] = []
        native_text = page.get_text("text").strip()
        if native_text:
            page_parts.append(native_text)

        tables, _ = extract_native_tables(input_path, page_number)
        if tables:
            camelot_pages += 1
            for table_index, frame in enumerate(tables, 1):
                table_md = dataframe_to_markdown(frame)
                if not table_md:
                    continue
                heading = f"### Table {table_index}" if len(tables) > 1 else "### Table"
                page_parts.append(f"{heading}\n\n{table_md}")
                table_count += 1

        add_page_heading(parts, page_number, total_pages, "\n\n".join(page_parts))

    doc.close()
    text = "\n\n".join(part for part in parts if part.strip()).strip()
    summary = build_summary("pdf", total_pages, scanned_pages, native_pages, table_count, ocr_fallback_pages, "offline pdf engine")

    return {
        "ok": True,
        "text": text,
        "meta": {
            "kind": "pdf",
            "pages": total_pages,
            "scannedPages": scanned_pages,
            "nativePages": native_pages,
            "tableCount": table_count,
            "ocrFallbackPages": ocr_fallback_pages,
            "camelotPages": camelot_pages,
            "modelDir": str(model_dir) if model_dir else "",
            "summary": summary,
        },
    }


def process_image(input_path: Path, model_dir: Path | None, threshold: float) -> dict:
    struct_engine, ocr_engine = get_engines(model_dir, threshold)
    image = cv2.imdecode(np.fromfile(str(input_path), dtype=np.uint8), cv2.IMREAD_COLOR)
    if image is None:
        raise RuntimeError(f"Unable to read image: {input_path}")

    text = ""
    table_count = 0
    ocr_fallback = False

    if struct_engine is not None:
        text, table_count, structured_chars = ppstructure_to_markdown(struct_engine(image))
    else:
        structured_chars = 0

    if ocr_engine is not None and structured_chars < 80:
        ocr_text, ocr_chars = ocr_to_markdown(ocr_engine.ocr(image, cls=True))
        if ocr_chars > structured_chars:
            text = ocr_text
            ocr_fallback = True

    if not text.strip():
        raise RuntimeError("No text could be extracted from the image")

    summary = build_summary("image", 1, 1, 0, table_count, 1 if ocr_fallback else 0, "offline pdf engine")
    return {
        "ok": True,
        "text": text.strip(),
        "meta": {
            "kind": "image",
            "pages": 1,
            "scannedPages": 1,
            "nativePages": 0,
            "tableCount": table_count,
            "ocrFallbackPages": 1 if ocr_fallback else 0,
            "camelotPages": 0,
            "modelDir": str(model_dir) if model_dir else "",
            "summary": summary,
        },
    }


def append_text_paragraphs(docx_doc: Document, text: str) -> int:
    count = 0
    for paragraph in [part.strip() for part in text.splitlines()]:
        if not paragraph:
            continue
        docx_doc.add_paragraph(paragraph)
        count += len(paragraph)
    return count


def append_dataframe_table(docx_doc: Document, df: pd.DataFrame) -> int:
    if df.empty and len(df.columns) == 0:
        return 0

    table = docx_doc.add_table(rows=df.shape[0] + 1, cols=max(df.shape[1], 1))
    table.style = "Table Grid"
    chars = 0

    for col_index, col_name in enumerate(df.columns.tolist()[: len(table.columns)]):
        table.cell(0, col_index).text = normalize_text(col_name)
        chars += len(normalize_text(col_name))

    for row_index in range(df.shape[0]):
        for col_index in range(df.shape[1]):
            cell_text = normalize_text(df.iloc[row_index, col_index])
            table.cell(row_index + 1, col_index).text = cell_text
            chars += len(cell_text)
    return chars


def append_ppstructure_to_docx(docx_doc: Document, blocks: object) -> tuple[int, int]:
    if not isinstance(blocks, list):
        return 0, 0

    chars = 0
    table_count = 0
    sorted_blocks = sorted(
        (block for block in blocks if isinstance(block, dict)),
        key=lambda block: (
            block.get("bbox", [0, 0, 0, 0])[1],
            block.get("bbox", [0, 0, 0, 0])[0],
        ),
    )

    for block in sorted_blocks:
        block_type = normalize_text(block.get("type")).lower()
        if block_type == "table":
            html = ""
            if isinstance(block.get("res"), dict):
                html = normalize_text(block["res"].get("html"))
            if not html:
                continue
            try:
                frames = pd.read_html(io.StringIO(html))
            except Exception:
                continue
            if not frames:
                continue
            chars += append_dataframe_table(docx_doc, frames[0].fillna(""))
            table_count += 1
            continue

        lines = extract_block_lines(block.get("res"))
        text = "\n".join(line for line in lines if normalize_text(line)).strip()
        if not text:
            continue
        if block_type == "title":
            docx_doc.add_heading(text, level=1)
        else:
            append_text_paragraphs(docx_doc, text)
        chars += len(text)
    return chars, table_count


def append_ocr_to_docx(docx_doc: Document, raw: object) -> int:
    text, chars = ocr_to_markdown(raw)
    append_text_paragraphs(docx_doc, text)
    return chars


def export_pdf_to_docx(input_path: Path, output_path: Path, model_dir: Path | None, threshold: float) -> dict:
    fitz_doc = fitz.open(str(input_path))
    total_pages = len(fitz_doc)
    scanned_flags = [is_scanned_page(fitz_doc.load_page(index)) for index in range(total_pages)]

    if total_pages > 0 and not any(scanned_flags):
        converter = Converter(str(input_path))
        try:
            converter.convert(str(output_path))
        finally:
            converter.close()
        fitz_doc.close()
        return {
            "ok": True,
            "meta": {
                "kind": "pdf",
                "pages": total_pages,
                "scannedPages": 0,
                "nativePages": total_pages,
                "tableCount": 0,
                "ocrFallbackPages": 0,
                "camelotPages": 0,
                "modelDir": str(model_dir) if model_dir else "",
                "summary": build_summary("pdf", total_pages, 0, total_pages, 0, 0, "pdf2docx direct"),
            },
        }

    struct_engine, ocr_engine = get_engines(model_dir, threshold)
    if any(scanned_flags) and struct_engine is None and ocr_engine is None:
        fitz_doc.close()
        raise RuntimeError("OCR models are unavailable for scanned PDF to DOCX export")

    docx_doc = Document()
    scanned_pages = 0
    native_pages = 0
    table_count = 0
    ocr_fallback_pages = 0

    for index in range(total_pages):
        page = fitz_doc.load_page(index)
        page_number = index + 1
        if total_pages > 1:
            docx_doc.add_heading(f"Page {page_number}", level=1)

        if scanned_flags[index]:
            scanned_pages += 1
            image = render_page_bgr(page)
            chars_written = 0
            if struct_engine is not None:
                page_chars, page_tables = append_ppstructure_to_docx(docx_doc, struct_engine(image))
                chars_written += page_chars
                table_count += page_tables

            if chars_written < 80 and ocr_engine is not None:
                ocr_chars = append_ocr_to_docx(docx_doc, ocr_engine.ocr(image, cls=True))
                if ocr_chars > chars_written:
                    chars_written = ocr_chars
                    ocr_fallback_pages += 1

            if chars_written == 0:
                fallback_text = page.get_text("text").strip()
                if fallback_text:
                    append_text_paragraphs(docx_doc, fallback_text)
        else:
            native_pages += 1
            native_text = page.get_text("text").strip()
            if native_text:
                append_text_paragraphs(docx_doc, native_text)

            tables, _ = extract_native_tables(input_path, page_number)
            for frame in tables:
                table_count += 1
                append_dataframe_table(docx_doc, frame)

        if index < total_pages - 1:
            docx_doc.add_page_break()

    fitz_doc.close()
    docx_doc.save(str(output_path))

    return {
        "ok": True,
        "meta": {
            "kind": "pdf",
            "pages": total_pages,
            "scannedPages": scanned_pages,
            "nativePages": native_pages,
            "tableCount": table_count,
            "ocrFallbackPages": ocr_fallback_pages,
            "camelotPages": 0,
            "modelDir": str(model_dir) if model_dir else "",
            "summary": build_summary("pdf", total_pages, scanned_pages, native_pages, table_count, ocr_fallback_pages, "structured docx export"),
        },
    }


def export_image_to_docx(input_path: Path, output_path: Path, model_dir: Path | None, threshold: float) -> dict:
    struct_engine, ocr_engine = get_engines(model_dir, threshold)
    image = cv2.imdecode(np.fromfile(str(input_path), dtype=np.uint8), cv2.IMREAD_COLOR)
    if image is None:
        raise RuntimeError(f"Unable to read image: {input_path}")

    docx_doc = Document()
    chars_written = 0
    table_count = 0
    ocr_fallback = 0

    if struct_engine is not None:
        chars_written, table_count = append_ppstructure_to_docx(docx_doc, struct_engine(image))

    if chars_written < 80 and ocr_engine is not None:
        ocr_chars = append_ocr_to_docx(docx_doc, ocr_engine.ocr(image, cls=True))
        if ocr_chars > chars_written:
            chars_written = ocr_chars
            ocr_fallback = 1

    if chars_written == 0:
        raise RuntimeError("No text could be extracted from the image for DOCX export")

    docx_doc.save(str(output_path))
    return {
        "ok": True,
        "meta": {
            "kind": "image",
            "pages": 1,
            "scannedPages": 1,
            "nativePages": 0,
            "tableCount": table_count,
            "ocrFallbackPages": ocr_fallback,
            "camelotPages": 0,
            "modelDir": str(model_dir) if model_dir else "",
            "summary": build_summary("image", 1, 1, 0, table_count, ocr_fallback, "structured docx export"),
        },
    }


def parse_args(argv: Iterable[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--action", choices=("extract", "export-docx"), default="extract")
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", default="")
    parser.add_argument("--type", choices=("pdf", "image"), required=True)
    parser.add_argument("--model-dir", default="")
    parser.add_argument("--threshold", type=float, default=0.5)
    return parser.parse_args(list(argv))


def main(argv: Iterable[str]) -> int:
    args = parse_args(argv)
    input_path = Path(args.input)
    output_path = Path(args.output) if args.output else None
    model_dir = Path(args.model_dir) if args.model_dir else None

    try:
        if args.action == "extract":
            if args.type == "pdf":
                payload = process_pdf(input_path, model_dir, args.threshold)
            else:
                payload = process_image(input_path, model_dir, args.threshold)
        else:
            if output_path is None:
                raise RuntimeError("output path is required for export-docx")
            if args.type == "pdf":
                payload = export_pdf_to_docx(input_path, output_path, model_dir, args.threshold)
            else:
                payload = export_image_to_docx(input_path, output_path, model_dir, args.threshold)

        sys.stdout.write(json.dumps(payload, ensure_ascii=False))
        return 0
    except Exception as exc:
        sys.stdout.write(
            json.dumps(
                {
                    "ok": False,
                    "error": str(exc),
                    "meta": {
                        "kind": args.type,
                        "modelDir": str(model_dir) if model_dir else "",
                    },
                },
                ensure_ascii=False,
            )
        )
        return 1


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
